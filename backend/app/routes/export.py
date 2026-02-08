"""Export routes for generating Word and PDF documents from markdown content"""
import io
import re
from urllib.parse import quote
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

# Word document generation
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

router = APIRouter(prefix="/export", tags=["export"])


class ExportRequest(BaseModel):
    content: str
    filename: str = "document"


def parse_markdown_to_blocks(md_content: str) -> list:
    """Parse markdown content into structured blocks for document generation"""
    blocks = []
    lines = md_content.split('\n')
    current_list = []
    in_list = False
    
    for line in lines:
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            if in_list and current_list:
                blocks.append({'type': 'list', 'items': current_list})
                current_list = []
                in_list = False
            continue
        
        # Headers
        if stripped.startswith('# '):
            if in_list and current_list:
                blocks.append({'type': 'list', 'items': current_list})
                current_list = []
                in_list = False
            blocks.append({'type': 'h1', 'text': stripped[2:]})
        elif stripped.startswith('## '):
            if in_list and current_list:
                blocks.append({'type': 'list', 'items': current_list})
                current_list = []
                in_list = False
            blocks.append({'type': 'h2', 'text': stripped[3:]})
        elif stripped.startswith('### '):
            if in_list and current_list:
                blocks.append({'type': 'list', 'items': current_list})
                current_list = []
                in_list = False
            blocks.append({'type': 'h3', 'text': stripped[4:]})
        elif stripped.startswith('#### '):
            if in_list and current_list:
                blocks.append({'type': 'list', 'items': current_list})
                current_list = []
                in_list = False
            blocks.append({'type': 'h4', 'text': stripped[5:]})
        # List items
        elif stripped.startswith('- ') or stripped.startswith('* '):
            in_list = True
            current_list.append(stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            in_list = True
            current_list.append(re.sub(r'^\d+\.\s', '', stripped))
        # Regular paragraph
        else:
            if in_list and current_list:
                blocks.append({'type': 'list', 'items': current_list})
                current_list = []
                in_list = False
            # Clean up markdown formatting
            text = stripped
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.+?)`', r'\1', text)  # Code
            blocks.append({'type': 'paragraph', 'text': text})
    
    # Don't forget remaining list
    if current_list:
        blocks.append({'type': 'list', 'items': current_list})
    
    return blocks


@router.post("/word")
async def export_to_word(request: ExportRequest):
    """Export markdown content to Word document"""
    try:
        doc = Document()
        
        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        blocks = parse_markdown_to_blocks(request.content)
        
        for block in blocks:
            if block['type'] == 'h1':
                p = doc.add_heading(block['text'], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif block['type'] == 'h2':
                p = doc.add_heading(block['text'], level=2)
            elif block['type'] == 'h3':
                p = doc.add_heading(block['text'], level=3)
            elif block['type'] == 'h4':
                p = doc.add_paragraph(block['text'])
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(12)
            elif block['type'] == 'list':
                for item in block['items']:
                    p = doc.add_paragraph(item, style='List Bullet')
            elif block['type'] == 'paragraph':
                p = doc.add_paragraph(block['text'])
                p.paragraph_format.space_after = Pt(6)
        
        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Properly encode filename for Content-Disposition header
        safe_filename = request.filename.replace(' ', '_')
        encoded_filename = quote(f"{safe_filename}.docx")
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating Word document: {str(e)}")


@router.post("/pdf")
async def export_to_pdf(request: ExportRequest):
    """Export markdown content to PDF document"""
    try:
        buffer = io.BytesIO()
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Register fonts with Cyrillic support
        try:
            pdfmetrics.registerFont(TTFont('DejaVuSans', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
            default_font = 'DejaVuSans'
            bold_font = 'DejaVuSans-Bold'
        except Exception:
            default_font = 'Helvetica'
            bold_font = 'Helvetica-Bold'
        
        # Create styles
        styles = getSampleStyleSheet()
        
        # Custom styles for better typography with Cyrillic support
        styles.add(ParagraphStyle(
            name='CustomTitle',
            fontName=bold_font,
            fontSize=18,
            spaceAfter=12,
            textColor=colors.HexColor('#1a1a1a')
        ))
        
        styles.add(ParagraphStyle(
            name='CustomH2',
            fontName=bold_font,
            fontSize=14,
            spaceBefore=12,
            spaceAfter=6,
            textColor=colors.HexColor('#333333')
        ))
        
        styles.add(ParagraphStyle(
            name='CustomH3',
            fontName=bold_font,
            fontSize=12,
            spaceBefore=10,
            spaceAfter=4,
            textColor=colors.HexColor('#444444')
        ))
        
        styles.add(ParagraphStyle(
            name='CustomBody',
            fontName=default_font,
            fontSize=10,
            spaceAfter=6,
            leading=14
        ))
        
        styles.add(ParagraphStyle(
            name='CustomListItem',
            fontName=default_font,
            fontSize=10,
            leftIndent=20,
            spaceAfter=3,
            bulletIndent=10
        ))
        
        story = []
        blocks = parse_markdown_to_blocks(request.content)
        
        for block in blocks:
            if block['type'] == 'h1':
                story.append(Paragraph(block['text'], styles['CustomTitle']))
            elif block['type'] == 'h2':
                story.append(Paragraph(block['text'], styles['CustomH2']))
            elif block['type'] == 'h3':
                story.append(Paragraph(block['text'], styles['CustomH3']))
            elif block['type'] == 'h4':
                story.append(Paragraph(f"<b>{block['text']}</b>", styles['CustomBody']))
            elif block['type'] == 'list':
                for item in block['items']:
                    story.append(Paragraph(f"• {item}", styles['CustomListItem']))
                story.append(Spacer(1, 6))
            elif block['type'] == 'paragraph':
                story.append(Paragraph(block['text'], styles['CustomBody']))
        
        doc.build(story)
        buffer.seek(0)
        
        # Properly encode filename for Content-Disposition header
        safe_filename = request.filename.replace(' ', '_')
        encoded_filename = quote(f"{safe_filename}.pdf")
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating PDF document: {str(e)}")
